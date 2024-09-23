import openai
from anthropic import Anthropic
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io
import streamlit as st
import json
import os
import base64
from docx2pdf import convert
import tempfile

# Chargement des clés API à partir de config.json
def load_api_keys():
    if os.path.exists('config.json'):
        with open('config.json', 'r') as f:
            return json.load(f)
    else:
        raise FileNotFoundError("Le fichier config.json est introuvable.")

keys = load_api_keys()

# Configuration des API
openai.api_key = keys['openai_key']
anthropic = Anthropic(api_key=keys['anthropic_key'])

class FormationManager:
    def __init__(self):
        self.document = Document()
        self.content = {}
        self.versions = {}
        self.current_version = 1
        self.load_memory()

    def load_memory(self):
        if os.path.exists('formation_memory.json'):
            with open('formation_memory.json', 'r') as f:
                data = json.load(f)
                self.content = data['content']
                self.versions = data['versions']
                self.current_version = data['current_version']

    def save_memory(self):
        with open('formation_memory.json', 'w') as f:
            json.dump({
                'content': self.content,
                'versions': self.versions,
                'current_version': self.current_version
            }, f)

    def generer_contenu_ia(self, prompt, section, use_openai=True):
        if use_openai:
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}]
            )
            content = response.choices[0].message.content
        else:
            response = anthropic.completions.create(
                model="claude-3-sonnet-20240229",
                prompt=prompt,
                max_tokens_to_sample=1000
            )
            content = response.completion
        self.content[section] = content
        return content

    def creer_graphique(self, donnees, titre):
        plt.figure(figsize=(10, 6))
        plt.bar(donnees.keys(), donnees.values())
        plt.title(titre)
        plt.xlabel("Catégories")
        plt.ylabel("Valeurs")
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png')
        img_buffer.seek(0)
        return img_buffer

    def ajouter_styles(self):
        styles = self.document.styles

        title_style = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.color.rgb = RGBColor(0, 0, 128)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        heading_style = styles.add_style('HeadingStyle', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = 'Calibri'
        heading_font.size = Pt(18)
        heading_font.color.rgb = RGBColor(0, 128, 0)

        body_style = styles.add_style('BodyStyle', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Georgia'
        body_font.size = Pt(11)

    def ajouter_table_des_matieres(self):
        self.document.add_paragraph("Table des matières", style='HeadingStyle')
        self.document.add_paragraph().add_run("").add_break()
        self.document.add_table_of_contents()
        self.document.add_page_break()

    def creer_formation(self):
        self.document = Document()
        self.ajouter_styles()

        self.document.add_heading('Formation de Coiffure à Domicile', 0).style = self.document.styles['TitleStyle']
        self.ajouter_table_des_matieres()

        sections = [
            ('introduction', "Écrivez une introduction pour une formation de coiffure à domicile."),
            ('techniques', "Listez et expliquez 5 techniques essentielles pour une coiffeuse à domicile."),
            ('equipement', "Détaillez l'équipement nécessaire pour une coiffeuse à domicile."),
            ('gestion_client', "Expliquez comment gérer efficacement une clientèle à domicile."),
            ('marketing', "Donnez des conseils de marketing pour une coiffeuse à domicile.")
        ]

        for section, prompt in sections:
            self.document.add_heading(section.capitalize(), level=1).style = self.document.styles['HeadingStyle']
            content = self.content.get(section) or self.generer_contenu_ia(prompt, section, use_openai=(section != 'techniques'))
            self.document.add_paragraph(content, style='BodyStyle')

        # Graphique
        donnees = {"Coupe": 30, "Coloration": 25, "Coiffage": 20, "Soins": 15, "Conseil": 10}
        img_buffer = self.creer_graphique(donnees, "Répartition des compétences")
        self.document.add_picture(img_buffer, width=Inches(6))

        self.save_memory()
        self.create_new_version()

    def modifier_section(self, section, nouveau_contenu):
        self.content[section] = nouveau_contenu
        self.save_memory()
        self.creer_formation()

    def create_new_version(self):
        self.current_version += 1
        self.versions[str(self.current_version)] = self.content.copy()
        self.save_memory()

    def revert_to_version(self, version):
        if version in self.versions:
            self.content = self.versions[version].copy()
            self.current_version = int(version)
            self.save_memory()
            self.creer_formation()

    def sauvegarder_document(self, nom_fichier):
        self.document.save(nom_fichier)

def docx_to_pdf(docx_path):
    pdf_path = docx_path.replace('.docx', '.pdf')
    convert(docx_path, pdf_path)
    return pdf_path

def get_binary_file_downloader_html(bin_file, file_label='File'):
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{os.path.basename(bin_file)}">Télécharger {file_label}</a>'
    return href

def main():
    st.title("Générateur et Gestionnaire Avancé de Formation de Coiffure à Domicile")

    manager = FormationManager()

    if st.button("Générer/Mettre à jour la formation"):
        manager.creer_formation()
        st.success("Formation générée avec succès!")

    sections = ['introduction', 'techniques', 'equipement', 'gestion_client', 'marketing']
    section = st.selectbox("Choisir une section à modifier", sections)
    nouveau_contenu = st.text_area("Nouveau contenu", value=manager.content.get(section, ""))
    
    if st.button("Modifier la section"):
        manager.modifier_section(section, nouveau_contenu)
        st.success(f"Section '{section}' mise à jour avec succès!")

    versions = list(manager.versions.keys())
    if versions:
        version = st.selectbox("Choisir une version à restaurer", versions)
        if st.button("Restaurer la version"):
            manager.revert_to_version(version)
            st.success(f"Version {version} restaurée avec succès!")

    if st.button("Sauvegarder et prévisualiser le document"):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            manager.sauvegarder_document(tmp_file.name)
            pdf_path = docx_to_pdf(tmp_file.name)
        
        st.markdown(get_binary_file_downloader_html(tmp_file.name, 'DOCX'), unsafe_allow_html=True)
        st.markdown(get_binary_file_downloader_html(pdf_path, 'PDF'), unsafe_allow_html=True)
        
        with open(pdf_path, "rb") as f:
            base64_pdf = base64.b64encode(f.read()).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
